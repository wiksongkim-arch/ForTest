"""在线与本地需求文档的统一读取能力。"""

from __future__ import annotations

import re
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Literal, Mapping, Protocol


RequirementSourceType = Literal["link", "file"]
SUPPORTED_LOCAL_DOCUMENT_SUFFIXES = frozenset(
    {".md", ".txt", ".docx", ".pdf", ".xlsx"}
)
MAX_LOCAL_DOCUMENT_BYTES = 50 * 1024 * 1024
MAX_EXPANDED_DOCUMENT_BYTES = 200 * 1024 * 1024
MAX_DOCUMENT_CHARACTERS = 2_000_000
MAX_SPREADSHEET_CELLS = 200_000


@dataclass(frozen=True)
class RequirementDocumentSource:
    """稳定描述一次生成任务所使用的需求来源。"""

    source_type: RequirementSourceType
    location: str

    @classmethod
    def create(
        cls,
        source_type: str,
        location: str | Path,
    ) -> "RequirementDocumentSource":
        normalized_type = str(source_type or "link").strip().casefold()
        if normalized_type not in {"link", "file"}:
            raise ValueError("需求文档来源类型无效")
        normalized_location = str(location or "").strip()
        if not normalized_location:
            raise ValueError("需求文档地址不能为空")
        if normalized_type == "file":
            normalized_location = str(
                validate_local_requirement_path(normalized_location)
            )
        return cls(normalized_type, normalized_location)  # type: ignore[arg-type]


@dataclass(frozen=True)
class RequirementDocument:
    """生成流程只消费这一种规范化文档，不感知具体平台。"""

    name: str
    text: str
    images: tuple[str, ...] = ()
    warnings: tuple[str, ...] = ()


class RequirementDocumentReader(Protocol):
    def read(self, source: RequirementDocumentSource) -> RequirementDocument: ...


class RequirementDocumentGateway:
    """按来源类型路由读取器，为后续接入其它在线文档 MCP 预留扩展点。"""

    def __init__(
        self,
        readers: Mapping[str, RequirementDocumentReader],
    ) -> None:
        self._readers = {
            str(key).strip().casefold(): value for key, value in readers.items()
        }

    def read(
        self,
        source: RequirementDocumentSource | str,
    ) -> RequirementDocument:
        normalized = (
            source
            if isinstance(source, RequirementDocumentSource)
            else RequirementDocumentSource.create("link", source)
        )
        reader = self._readers.get(normalized.source_type)
        if reader is None:
            raise RuntimeError("当前需求文档来源尚未接入读取能力")
        document = reader.read(normalized)
        if not isinstance(document, RequirementDocument):
            raise RuntimeError("需求文档读取结果格式无效")
        text = str(document.text or "").strip()
        if not text:
            raise RuntimeError("需求文档正文为空")
        if len(text) > MAX_DOCUMENT_CHARACTERS:
            raise RuntimeError("需求文档正文过大")
        name = str(document.name or "").strip() or "测试用例"
        return RequirementDocument(
            name=name[:200],
            text=text,
            images=tuple(str(item) for item in document.images if str(item)),
            warnings=tuple(str(item) for item in document.warnings if str(item)),
        )


class DingTalkRequirementDocumentReader:
    """把既有钉钉 MCP 返回值转换成统一文档。"""

    def __init__(self, document_service: Any) -> None:
        self.document_service = document_service

    def read(self, source: RequirementDocumentSource) -> RequirementDocument:
        if source.source_type != "link":
            raise ValueError("钉钉读取器只接受链接来源")
        warnings: list[str] = []
        try:
            raw_name = self.document_service.get_document_name(source.location)
            name = str(raw_name or "").strip() or "测试用例"
        except Exception:
            name = "测试用例"
            warnings.append("获取文档名称失败，使用默认名称")

        payload = self.document_service.get_document_content(source.location)
        if not isinstance(payload, dict) or payload.get("isError") is True:
            raise RuntimeError("文档内容不可用")
        text = _payload_text(payload)
        supplied_images = payload.get("images", ())
        images = (
            tuple(item for item in supplied_images if isinstance(item, str))
            if isinstance(supplied_images, (list, tuple))
            else ()
        )
        return RequirementDocument(
            name=name,
            text=text,
            images=images,
            warnings=tuple(warnings),
        )


class LocalRequirementDocumentReader:
    """安全读取用户明确选择的本地需求文档。"""

    def read(self, source: RequirementDocumentSource) -> RequirementDocument:
        if source.source_type != "file":
            raise ValueError("本地读取器只接受文件来源")
        path = validate_local_requirement_path(source.location)
        suffix = path.suffix.casefold()
        if suffix in {".md", ".txt"}:
            text = self._read_text(path)
        elif suffix == ".docx":
            self._validate_zip_expansion(path)
            text = self._read_docx(path)
        elif suffix == ".pdf":
            text = self._read_pdf(path)
        elif suffix == ".xlsx":
            self._validate_zip_expansion(path)
            text = self._read_xlsx(path)
        else:  # pragma: no cover - 路径验证已负责拦截
            raise ValueError("不支持的本地需求文档格式")
        normalized = str(text or "").strip()
        if not normalized:
            raise RuntimeError("本地需求文档正文为空")
        if len(normalized) > MAX_DOCUMENT_CHARACTERS:
            raise RuntimeError("本地需求文档正文过大")
        return RequirementDocument(name=path.stem, text=normalized)

    @staticmethod
    def _read_text(path: Path) -> str:
        raw = path.read_bytes()
        for encoding in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                # 二进制读取不会执行通用换行转换，这里统一成生成器使用的 LF。
                return raw.decode(encoding).replace("\r\n", "\n").replace(
                    "\r", "\n"
                )
            except UnicodeDecodeError:
                continue
        raise RuntimeError("本地文本文件编码不受支持")

    @staticmethod
    def _read_docx(path: Path) -> str:
        # 重型解析库延迟导入，普通在线任务和程序启动不会承担加载成本。
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        document = Document(str(path))
        lines: list[str] = []
        for child in document.element.body.iterchildren():
            if child.tag.endswith("}p"):
                paragraph = Paragraph(child, document)
                value = paragraph.text.strip()
                if not value:
                    continue
                style_name = str(getattr(paragraph.style, "name", "") or "")
                heading = re.search(r"(?:heading|标题)\s*([1-6])", style_name, re.I)
                lines.append(
                    f"{'#' * int(heading.group(1))} {value}" if heading else value
                )
            elif child.tag.endswith("}tbl"):
                table = Table(child, document)
                for row in table.rows:
                    values = [
                        cell.text.replace("|", "\\|").strip() for cell in row.cells
                    ]
                    if any(values):
                        lines.append("| " + " | ".join(values) + " |")
        return "\n".join(lines)

    @staticmethod
    def _read_pdf(path: Path) -> str:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        if reader.is_encrypted and reader.decrypt("") == 0:
            raise RuntimeError("本地 PDF 已加密，无法读取")
        if len(reader.pages) > 1000:
            raise RuntimeError("本地 PDF 页数过多")
        pages: list[str] = []
        for index, page in enumerate(reader.pages, start=1):
            value = str(page.extract_text() or "").strip()
            if value:
                pages.append(f"# 第 {index} 页\n{value}")
        return "\n\n".join(pages)

    @staticmethod
    def _read_xlsx(path: Path) -> str:
        from openpyxl import load_workbook

        workbook = load_workbook(path, read_only=True, data_only=True)
        sections: list[str] = []
        try:
            total_cells = 0
            for sheet in workbook.worksheets:
                # 兼容部分在线表格导出文件不准确的 dimension 元数据。
                reset_dimensions = getattr(sheet, "reset_dimensions", None)
                if callable(reset_dimensions):
                    reset_dimensions()
                rows: list[str] = []
                for values in sheet.iter_rows(values_only=True):
                    total_cells += len(values)
                    if total_cells > MAX_SPREADSHEET_CELLS:
                        raise RuntimeError("本地表格有效区域过大")
                    rendered = [
                        str(value).replace("|", "\\|").strip()
                        if value is not None
                        else ""
                        for value in values
                    ]
                    while rendered and not rendered[-1]:
                        rendered.pop()
                    if rendered and any(rendered):
                        rows.append("| " + " | ".join(rendered) + " |")
                if rows:
                    sections.append(f"# {sheet.title}\n" + "\n".join(rows))
        finally:
            workbook.close()
        return "\n\n".join(sections)

    @staticmethod
    def _validate_zip_expansion(path: Path) -> None:
        try:
            with zipfile.ZipFile(path) as archive:
                members = archive.infolist()
                if len(members) > 10_000:
                    raise RuntimeError("本地文档压缩包条目过多")
                expanded = sum(max(0, item.file_size) for item in members)
                if expanded > MAX_EXPANDED_DOCUMENT_BYTES:
                    raise RuntimeError("本地文档解压后体积过大")
        except zipfile.BadZipFile:
            raise RuntimeError("本地 Office 文档格式无效") from None


def validate_local_requirement_path(value: str | Path) -> Path:
    """返回已校验的绝对文件路径，错误信息不携带用户目录。"""

    raw = str(value or "").strip()
    if not raw:
        raise ValueError("请选择本地需求文档")
    path = Path(raw).expanduser()
    if not path.is_absolute():
        raise ValueError("本地需求文档必须使用绝对路径")
    try:
        path = path.resolve(strict=True)
    except OSError:
        raise ValueError("本地需求文档不存在") from None
    if not path.is_file():
        raise ValueError("本地需求文档不存在")
    if path.suffix.casefold() not in SUPPORTED_LOCAL_DOCUMENT_SUFFIXES:
        raise ValueError("不支持的本地需求文档格式")
    size = path.stat().st_size
    if size <= 0:
        raise ValueError("本地需求文档为空")
    if size > MAX_LOCAL_DOCUMENT_BYTES:
        raise ValueError("本地需求文档不能超过 50 MB")
    return path


def _payload_text(payload: Mapping[str, Any]) -> str:
    text_value = payload.get(
        "texts",
        payload.get("markdown", payload.get("text", "")),
    )
    if isinstance(text_value, list):
        text = "\n".join(item for item in text_value if isinstance(item, str))
    elif isinstance(text_value, str):
        text = text_value
    else:
        raise RuntimeError("文档正文格式无效")
    if not text.strip():
        raise RuntimeError("文档正文为空")
    return text


__all__ = [
    "DingTalkRequirementDocumentReader",
    "LocalRequirementDocumentReader",
    "RequirementDocument",
    "RequirementDocumentGateway",
    "RequirementDocumentReader",
    "RequirementDocumentSource",
    "SUPPORTED_LOCAL_DOCUMENT_SUFFIXES",
    "validate_local_requirement_path",
]
