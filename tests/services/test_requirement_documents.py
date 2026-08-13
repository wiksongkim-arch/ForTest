"""本地与在线需求文档统一读取能力测试。"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from openpyxl import Workbook

from services.requirement_documents import (
    DingTalkRequirementDocumentReader,
    LocalRequirementDocumentReader,
    RequirementDocumentGateway,
    RequirementDocumentSource,
)


class _FakeDingTalkDocument:
    def get_document_name(self, _location: str) -> str:
        return "在线需求"

    def get_document_content(self, _location: str) -> dict:
        return {
            "texts": ["# 登录", "支持账号密码登录"],
            "images": ["https://example.test/login.png"],
        }


def test_gateway_routes_online_document_through_the_same_read_contract():
    gateway = RequirementDocumentGateway(
        {"link": DingTalkRequirementDocumentReader(_FakeDingTalkDocument())}
    )

    document = gateway.read(
        RequirementDocumentSource.create(
            "link",
            "https://alidocs.dingtalk.com/i/nodes/example",
        )
    )

    assert document.name == "在线需求"
    assert "支持账号密码登录" in document.text
    assert document.images == ("https://example.test/login.png",)


@pytest.mark.parametrize("suffix", [".md", ".txt"])
def test_local_text_document_uses_absolute_path_and_keeps_headings(
    tmp_path: Path,
    suffix: str,
):
    source = tmp_path / f"登录需求{suffix}"
    source.write_text("# 登录\n支持验证码登录", encoding="utf-8")
    gateway = RequirementDocumentGateway(
        {"file": LocalRequirementDocumentReader()}
    )

    document = gateway.read(
        RequirementDocumentSource.create("file", source.resolve())
    )

    assert document.name == "登录需求"
    assert document.text == "# 登录\n支持验证码登录"


def test_local_docx_preserves_heading_and_table_content(tmp_path: Path):
    source = tmp_path / "订单需求.docx"
    value = Document()
    value.add_heading("订单列表", level=1)
    value.add_paragraph("支持按订单号搜索")
    table = value.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "状态"
    table.cell(0, 1).text = "已完成"
    value.save(source)

    document = LocalRequirementDocumentReader().read(
        RequirementDocumentSource.create("file", source.resolve())
    )

    assert "# 订单列表" in document.text
    assert "支持按订单号搜索" in document.text
    assert "| 状态 | 已完成 |" in document.text


def test_local_xlsx_groups_content_by_sheet(tmp_path: Path):
    source = tmp_path / "表格需求.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "退款"
    sheet.append(["功能", "说明"])
    sheet.append(["申请退款", "必须填写原因"])
    workbook.save(source)

    document = LocalRequirementDocumentReader().read(
        RequirementDocumentSource.create("file", source.resolve())
    )

    assert "# 退款" in document.text
    assert "| 申请退款 | 必须填写原因 |" in document.text


def test_local_document_rejects_relative_or_unsupported_files(tmp_path: Path):
    unsupported = tmp_path / "需求.exe"
    unsupported.write_bytes(b"not-a-document")

    with pytest.raises(ValueError, match="绝对路径"):
        RequirementDocumentSource.create("file", "需求.md")
    with pytest.raises(ValueError, match="不支持"):
        RequirementDocumentSource.create("file", unsupported.resolve())
